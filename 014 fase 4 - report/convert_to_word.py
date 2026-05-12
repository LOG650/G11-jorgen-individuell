"""
Convert Final report.md to Final report.docx using pandoc (via pypandoc-binary).

Strategy:
- Strip the leading H1 + LOG650/group markup from the markdown body and pass
  these as pandoc metadata instead, so the docx has a proper title page without
  duplicating the title.
- Keep TOC, numbered sections, and the figures/ directory resolvable from cwd.

Run:  python convert_to_word.py
"""

from pathlib import Path
import re
import pypandoc

BASE = Path(__file__).parent
SRC = BASE / "Final report.md"
DST = BASE / "Final report.docx"

raw = SRC.read_text(encoding="utf-8")

# Strip leading H1 + author/course lines (we pass them as pandoc metadata).
# Match the first H1 and the two bold-line author/course markers right after it,
# then the first --- separator. Replace with empty so the body starts at
# "## Sammendrag" or the status blockquote.
pattern = re.compile(
    r"^# NautiCost:.*?\n+"          # H1 title
    r"(\*\*LOG650.*?\*\*\n+)"        # course line
    r"(\*\*Gruppe.*?\*\*\n+)",       # group line
    re.DOTALL | re.MULTILINE,
)
body = pattern.sub("", raw, count=1)

# Write the slimmed-down markdown to a temp file pandoc reads
TMP = BASE / "_tmp_for_pandoc.md"
TMP.write_text(body, encoding="utf-8")

try:
    pypandoc.convert_file(
        str(TMP),
        "docx",
        outputfile=str(DST),
        extra_args=[
            "--standalone",
            "--toc",
            "--toc-depth=3",
            "--number-sections",
            f"--resource-path={BASE}",
            "--metadata=title:NautiCost — Datadreven kostnadsestimering for yachthavneanløp i Skandinavia",
            "--metadata=subtitle:LOG650 Forskningsprosjekt, vår 2026 — Gruppe 11, Jørgen Renè (individuell)",
            "--metadata=author:Jørgen Renè",
            "--metadata=date:2026-05-13",
        ],
    )
finally:
    TMP.unlink(missing_ok=True)

# Sanity check
from docx import Document
doc = Document(DST)
print(f"Skrev {DST}")
print(f"  Avsnitt: {len(doc.paragraphs)}")
print(f"  Tabeller: {len(doc.tables)}")
print(f"  Bilder: {len(doc.inline_shapes)}")
print(f"  Størrelse: {DST.stat().st_size // 1024} KB")
print()
print("Første 6 elementer (style + tekst):")
for p in doc.paragraphs[:8]:
    style = p.style.name
    text = p.text[:80]
    if text:
        print(f"  [{style}] {text}")
