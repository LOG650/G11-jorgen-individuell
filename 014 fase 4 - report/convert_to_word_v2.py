"""
Convert "Final report draft.md" -> .docx -> .pdf

Forsiden bygges direkte i markdown med HiM-logoen (figures/him-logo.png),
tittel, subtittel og bibliografiske data. Post-prosessen sentrerer alle
forside-elementer, gjør tittelen til en stor sentrert overskrift og setter
luftig topp-marg på logoen.

Body etter forsiden følger standard Word-stiler:
- Times New Roman 12 pt med 1,5 linjeavstand
- Overskrift 1/2/3 fra Word-malen (Calibri Light, mørkeblå)

TOC og PDF eksporteres via Word COM i samme sesjon.

Run:  python convert_to_word_v2.py
"""

from pathlib import Path
import re
import shutil
import tempfile
import pypandoc
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent
SRC = BASE / "Final report draft.md"
DST_DOCX = BASE / "Final report draft.docx"
DST_PDF = BASE / "Final report draft.pdf"

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
LINE_SPACING = 1.5

HEADING_STYLES = {
    "Title", "Subtitle",
    "Heading 1", "Heading 2", "Heading 3",
    "Heading 4", "Heading 5", "Heading 6",
    "TOC Heading",
}

# 1) Pandoc → docx
raw = SRC.read_text(encoding="utf-8")
TMP = BASE / "_tmp_for_pandoc_v2.md"
TMP.write_text(raw, encoding="utf-8")

try:
    pypandoc.convert_file(
        str(TMP),
        "docx",
        outputfile=str(DST_DOCX),
        extra_args=[
            "--standalone",
            f"--resource-path={BASE}",
        ],
    )
finally:
    TMP.unlink(missing_ok=True)


# 2) Stiler: body-paragrafer får TNR 12 pt + 1,5 linjeavstand;
#    overskrifter beholdes som Word-stilene definerer dem.
def _apply_run_font(run, size=None):
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    from docx.oxml.ns import qn
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_paragraph_spacing(paragraph, line_spacing=LINE_SPACING):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


doc = Document(DST_DOCX)

# Oppdater body-stiler
body_styles = ["Normal", "Body Text", "First Paragraph", "Compact",
               "Caption", "Image Caption", "Quote", "Intense Quote",
               "List Paragraph", "Block Text"]
for sname in body_styles:
    try:
        s = doc.styles[sname]
        s.font.name = FONT_NAME
        if sname not in ("Caption", "Image Caption"):
            s.font.size = BODY_SIZE
    except KeyError:
        pass


# 3) Forside-styling: sentrer logo, tittel, subtittel og bibliografiske
#    linjer. Gjør tittelen visuelt stor (48 pt) og legg på topp-marg
#    over logoen slik at innholdet sitter pent på siden.
COVER_INDEX_END = None
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.style.name == "Heading 2":
        # Første Overskrift 2 = "Obligatorisk egenerklæring" → utenfor forside
        COVER_INDEX_END = i
        break

if COVER_INDEX_END is None:
    COVER_INDEX_END = 0  # ingen H2 → ingen forside å style

for i, paragraph in enumerate(doc.paragraphs[:COVER_INDEX_END]):
    # Sentrer alt på forsiden
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # Sjekk om paragrafen inneholder logoen (bildet)
    from docx.oxml.ns import qn
    has_image = bool(paragraph._element.findall('.//' + qn('w:drawing')))
    if has_image:
        # Gi luftig topp-marg over logoen (≈ 3 cm)
        paragraph.paragraph_format.space_before = Pt(72)
        paragraph.paragraph_format.space_after = Pt(48)

    # Sjekk om paragrafen er tittelen (Heading 1)
    if paragraph.style.name == "Heading 1":
        paragraph.paragraph_format.space_before = Pt(36)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in paragraph.runs:
            run.font.size = Pt(48)
            run.font.bold = True
    elif paragraph.style.name in ("First Paragraph", "Body Text", "Normal"):
        for run in paragraph.runs:
            _apply_run_font(run, size=Pt(14) if i == COVER_INDEX_END - 1 else None)
        # Bruk større skrift på subtittel (linje rett etter Heading 1)
        # Identifiser subtittel: italic body etter tittel
        if i > 0:
            prev = doc.paragraphs[i - 1]
            if prev.style.name == "Heading 1":
                for run in paragraph.runs:
                    run.font.size = Pt(16)
                    run.font.italic = True

# Body-paragrafer (etter forsiden): TNR 12 pt + 1,5 linjeavstand
for paragraph in doc.paragraphs[COVER_INDEX_END:]:
    if paragraph.style.name in HEADING_STYLES:
        continue
    _set_paragraph_spacing(paragraph)
    for run in paragraph.runs:
        _apply_run_font(run, size=BODY_SIZE)

# Tabeller
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_spacing(paragraph)
                for run in paragraph.runs:
                    _apply_run_font(run, size=BODY_SIZE)

doc.save(DST_DOCX)
print(f"Stiler oppdatert i {DST_DOCX.name}")


# 4) TOC + PDF via Word COM i én sesjon, fra ASCII-temp-mappe
print()
print("Oppdaterer TOC og eksporterer PDF via Word COM …")
import win32com.client
import win32com.client.dynamic
import pythoncom

tmpdir = tempfile.mkdtemp(prefix="nauticost_")
tmp_docx = Path(tmpdir) / "report.docx"
tmp_pdf = Path(tmpdir) / "report.pdf"
shutil.copy2(DST_DOCX, tmp_docx)

try:
    pythoncom.CoInitialize()
    # Dynamisk (late) binding unngår en korrupt win32com gen_py-cache som
    # ellers gir "CLSIDToClassMap"/"property can not be set"-feil.
    word_app = win32com.client.dynamic.Dispatch("Word.Application")
    word_app.Visible = False
    word_app.DisplayAlerts = False
    wdoc = word_app.Documents.Open(str(tmp_docx))

    # Sidetall: sentrert PAGE-felt i bunntekst på alle seksjoner.
    # DifferentFirstPage gjør at forsiden (side 1) står uten tall.
    # Isolert i egen try/except slik at en evt. feil her ikke stopper
    # TOC-oppdatering og PDF-eksport.
    WD_HEADER_FOOTER_PRIMARY = 1
    WD_FIELD_PAGE = 33
    WD_ALIGN_PARAGRAPH_CENTER = 1
    try:
        for section in wdoc.Sections:
            footer = section.Footers(WD_HEADER_FOOTER_PRIMARY)
            footer.LinkToPrevious = False
            rng = footer.Range
            rng.Text = ""
            rng.ParagraphFormat.Alignment = WD_ALIGN_PARAGRAPH_CENTER
            wdoc.Fields.Add(rng, WD_FIELD_PAGE)
        # Forsiden (side 1) uten tall – best effort, riktig property­navn er
        # DifferentFirstPageHeaderFooter (ikke DifferentFirstPage).
        try:
            for section in wdoc.Sections:
                section.PageSetup.DifferentFirstPageHeaderFooter = True
            print("  Sidetall lagt til i bunntekst (forside uten tall).")
        except Exception as e2:
            print(f"  Sidetall lagt til (også på forsiden; "
                  f"forside-unntak feilet: {e2}).")
    except Exception as e:
        print(f"  Advarsel: kunne ikke legge til sidetall ({e}).")

    for toc in wdoc.TablesOfContents:
        toc.Update()
    wdoc.Fields.Update()
    wdoc.Save()
    wdoc.SaveAs2(str(tmp_pdf), FileFormat=17)
    wdoc.Close(False)
    word_app.Quit()
    pythoncom.CoUninitialize()
    shutil.copy2(tmp_docx, DST_DOCX)
    shutil.copy2(tmp_pdf, DST_PDF)
    print(f"  TOC oppdatert. PDF: {DST_PDF.name} ({DST_PDF.stat().st_size // 1024} KB)")

    # Post-fix: rename "Table of Contents" → "Innholdsfortegnelse"
    import zipfile
    tmp_zip = str(DST_DOCX) + ".tmp"
    with zipfile.ZipFile(DST_DOCX, "r") as zin:
        with zipfile.ZipFile(tmp_zip, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = data.replace(b"Table of Contents", b"Innholdsfortegnelse")
                zout.writestr(item, data)
    shutil.move(tmp_zip, str(DST_DOCX))
except Exception as e:
    print(f"  Word COM feilet: {e}")
finally:
    shutil.rmtree(tmpdir, ignore_errors=True)


# 5) Sanity check
doc = Document(DST_DOCX)
print()
print(f"Skrev {DST_DOCX.name}")
print(f"  Avsnitt: {len(doc.paragraphs)}")
print(f"  Tabeller: {len(doc.tables)}")
print(f"  Bilder: {len(doc.inline_shapes)}")
print(f"  Størrelse: {DST_DOCX.stat().st_size // 1024} KB")
